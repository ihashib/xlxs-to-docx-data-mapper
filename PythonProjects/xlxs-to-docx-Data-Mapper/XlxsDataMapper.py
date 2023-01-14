import pandas as pd
from docx import Document
from docx.shared import Inches
import json

#document to be printed
document = Document()

#convert .xlxs file to JSON file
excel_data_df = pd.read_excel('data.xlsx')
excel_data_df.to_json (r'Data.json')

#Use the JSON file to extract info
df = pd.DataFrame(pd.read_json("Data.json"))

f_name = df['Full Name']
organization = df['Organization']
title = df['Title']
phone = df['Phone']
email = df['Email']
addr = df['Full Address']

#Insert data on doc
size = len(df)
for i in range(size):
    document.add_paragraph('To, ')
    document.add_paragraph(f_name[i])
    document.add_paragraph(title[i])
    document.add_paragraph(organization[i])
    document.add_paragraph(addr[i])
    document.add_paragraph(str(phone[i]))
    document.add_paragraph(email[i])

    document.add_paragraph()

document.save('FinalDemo.docx')







